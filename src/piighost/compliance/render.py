"""Render structured compliance docs (ProcessingRegister / DPIAScreening /
SubjectAccessReport) to Markdown / DOCX / PDF deliverables.

Template resolution chain (first hit wins):
  1. User override:    ~/.piighost/templates/<profile>/<doctype>.<format>.j2
  2. Profile bundled:  piighost.compliance.templates/<profile>/<doctype>.<format>.j2
  3. Generic bundled:  piighost.compliance.templates/generic/<doctype>.<format>.j2

Markdown rendering only requires Jinja2 (declared in the [compliance] extra).
DOCX rendering additionally requires python-docx; PDF requires weasyprint
plus markdown. Missing optional deps raise a clear ImportError at call time
rather than at import.

Privacy invariant: the renderer consumes ONLY the ``data`` dict passed in
(already a sanitized Pydantic dump from the compliance subsystem). It does
not perform any secondary lookups, filesystem reads of source documents,
or vault accesses. Templates therefore cannot leak raw PII.
"""
from __future__ import annotations

import re
import time
from importlib import resources
from pathlib import Path
from typing import Any, Literal, Union

from pydantic import TypeAdapter, ValidationError

from piighost.service.models import (
    DPIAScreening,
    ProcessingRegister,
    SubjectAccessReport,
)

# Profile names are user-controlled and end up concatenated into a Path.
# Restrict to a strict identifier shape to block traversal / weird input
# before any filesystem lookup.
_PROFILE_RE = re.compile(r"^[a-z][a-z0-9_]{0,31}$")

_COMPLIANCE_UNION_ADAPTER = TypeAdapter(
    Union[ProcessingRegister, DPIAScreening, SubjectAccessReport]
)


def _validate_compliance_dict(data: dict) -> None:
    """Raise ValueError if *data* doesn't match any known compliance model.

    Uses ``model_config.extra = "forbid"`` semantics through the union
    adapter so unknown keys are rejected — closes the poisoned-dict
    attack surface flagged in Phase 2 followup #3.
    """
    if not isinstance(data, dict):
        raise ValueError(
            f"render_compliance_doc data must be a dict; got {type(data).__name__}"
        )
    try:
        _COMPLIANCE_UNION_ADAPTER.validate_python(data)
    except ValidationError as exc:
        raise ValueError(
            "render_compliance_doc data does not match any known compliance "
            f"model (ProcessingRegister / DPIAScreening / SubjectAccessReport): {exc}"
        ) from exc

# Doctype detection: which top-level keys identify which structured doc?
_DOCTYPE_MARKERS = {
    "registre": {"controller", "processing_name", "data_categories"},
    "dpia_screening": {"verdict", "triggers", "cnil_pia_inputs"},
    "subject_access": {"subject_tokens", "categories_found", "documents"},
}


def _detect_doctype(data: dict) -> str:
    """Inspect dict keys to choose a template family."""
    keys = set(data.keys())
    # Pick the marker set with the largest intersection.
    best: tuple[str, int] = ("registre", -1)
    for doctype, markers in _DOCTYPE_MARKERS.items():
        score = len(markers & keys)
        if score > best[1]:
            best = (doctype, score)
    if best[1] <= 0:
        # No marker matched — fall back to generic registre.
        return "registre"
    return best[0]


def _user_template_root() -> Path:
    return Path.home() / ".piighost" / "templates"


def _load_template(
    profile: str, doctype: str, fmt: str,
) -> tuple[str, str]:
    """Resolve the template chain and return ``(name, source)``.

    The returned ``name`` is a stable relative identifier (e.g.
    ``"generic/registre.md.j2"``) used by Jinja for include/inheritance
    bookkeeping; ``source`` is the template body.

    Raises FileNotFoundError if no template can be located.
    """
    if not _PROFILE_RE.match(profile):
        raise ValueError(f"Invalid profile name: {profile!r}")

    rel = f"{doctype}.{fmt}.j2"

    # 1. user override
    user_path = _user_template_root() / profile / rel
    if user_path.is_file():
        return (
            f"{profile}/{rel}",
            user_path.read_text(encoding="utf-8"),
        )

    # 2. bundled profile-specific
    pkg_root = resources.files("piighost.compliance.templates")
    profile_res = pkg_root / profile / rel
    if profile_res.is_file():
        return (
            f"{profile}/{rel}",
            profile_res.read_text(encoding="utf-8"),
        )

    # 3. fallback to bundled generic
    generic_res = pkg_root / "generic" / rel
    if generic_res.is_file():
        return (
            f"generic/{rel}",
            generic_res.read_text(encoding="utf-8"),
        )

    raise FileNotFoundError(
        f"No template found for profile={profile!r} doctype={doctype!r} "
        f"format={fmt!r}. Looked at user override, bundled profile, "
        f"and bundled generic locations.",
    )


def _build_jinja_env():
    """Construct a Jinja2 Environment that supports include/extends across
    the bundled-template tree AND user overrides.

    A FunctionLoader is used so the same fallback chain that resolves the
    main template also resolves any ``{% include %}`` directives inside it.
    """
    try:
        from jinja2 import Environment, FunctionLoader, select_autoescape
    except ImportError as e:
        raise ImportError(
            "Jinja2 is required for compliance rendering. "
            "Install with: pip install 'piighost[compliance]'",
        ) from e

    pkg_root = resources.files("piighost.compliance.templates")
    user_root = _user_template_root()

    def _load(name: str) -> str | None:
        # name is a relative path like "generic/registre.md.j2" or
        # "avocat/registre.md.j2". Try user override first, then bundled.
        user_path = user_root / name
        if user_path.is_file():
            return user_path.read_text(encoding="utf-8")
        bundled = pkg_root / name
        if bundled.is_file():
            return bundled.read_text(encoding="utf-8")
        return None

    env = Environment(
        loader=FunctionLoader(_load),
        autoescape=select_autoescape(disabled_extensions=("j2", "md", "j2.md")),
        keep_trailing_newline=True,
        trim_blocks=False,
        lstrip_blocks=False,
    )
    return env


def _render_md(
    data: dict, profile: str, doctype: str,
) -> str:
    """Render the Markdown body for the given data + profile."""
    env = _build_jinja_env()
    # Resolve the root template name so include/extends inside it can
    # use the same loader.
    name, _source = _load_template(profile, doctype, "md")
    template = env.get_template(name)
    return template.render(**data)


def _render_pdf(md_body: str, output_path: Path) -> int:
    """Convert Markdown -> HTML -> PDF via markdown + weasyprint."""
    try:
        import markdown as md_lib  # type: ignore[import-not-found]
    except ImportError as e:
        raise ImportError(
            "The 'markdown' package is required for PDF rendering. "
            "Install with: pip install 'piighost[compliance]'",
        ) from e
    try:
        from weasyprint import HTML  # type: ignore[import-not-found]
    except ImportError as e:
        raise ImportError(
            "weasyprint is required for PDF rendering. Install it "
            "manually (it is intentionally NOT in [compliance] because "
            "it pulls heavy native deps).",
        ) from e

    html_body = md_lib.markdown(
        md_body, extensions=["tables", "fenced_code"],
    )
    html_doc = (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        "<style>body{font-family:sans-serif;max-width:800px;margin:2em auto;}"
        "table{border-collapse:collapse;}"
        "th,td{border:1px solid #ccc;padding:4px 8px;}"
        "</style></head><body>"
        f"{html_body}"
        "</body></html>"
    )
    HTML(string=html_doc).write_pdf(str(output_path))
    return output_path.stat().st_size


def _render_docx(md_body: str, output_path: Path) -> int:
    """Render a minimal DOCX from the Markdown body.

    This is intentionally low-fi — we walk the Markdown line-by-line and
    emit headings / paragraphs. Tables and lists become plain paragraphs.
    For a richer pipeline plug docxtpl with a dedicated .docx template.
    """
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX rendering. Install it "
            "manually (it is intentionally NOT in [compliance] until "
            "DOCX output is officially supported).",
        ) from e

    doc = Document()
    for raw_line in md_body.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style=None)
        else:
            doc.add_paragraph(line)
    doc.save(str(output_path))
    return output_path.stat().st_size


def render_compliance_doc(
    *,
    data: dict[str, Any],
    format: Literal["md", "docx", "pdf"] = "md",
    profile: str = "generic",
    output_path: str | None = None,
) -> dict[str, Any]:
    """Render ``data`` (a Pydantic ``model_dump()``) to disk.

    Parameters
    ----------
    data:
        Sanitized dict from ProcessingRegister / DPIAScreening /
        SubjectAccessReport. NEVER pass raw vault entries here — the
        compliance models guarantee no PII leaks.
    format:
        One of ``md`` / ``docx`` / ``pdf``. PDF and DOCX require optional
        dependencies; missing deps raise a clear ImportError.
    profile:
        Template profile (``generic``, ``avocat``, ``notaire``,
        ``medecin``, ``expert_comptable``, ``rh``). Falls back to
        ``generic`` if a profession-specific template is missing.
    output_path:
        Destination path. If omitted, defaults to
        ``~/.piighost/exports/<project>-<doctype>-<ts>.<ext>``. Must
        resolve under ``~/.piighost/`` — the renderer rejects paths
        outside that root to keep MCP-callable surfaces from clobbering
        arbitrary files.

    Returns
    -------
    dict
        ``{path, format, size_bytes, rendered_at}`` — matches the shape
        of :class:`piighost.service.models.RenderResult`.
    """
    # Validate data against the 3-way union BEFORE doing any I/O.
    # This blocks adversarial input from a poisoned MCP context.
    _validate_compliance_dict(data)

    if format not in ("md", "docx", "pdf"):
        raise ValueError(f"Unsupported format: {format!r}")
    if not _PROFILE_RE.match(profile):
        raise ValueError(f"Invalid profile name: {profile!r}")

    doctype = _detect_doctype(data)

    if not output_path:
        project = data.get("project", "doc")
        ts = int(time.time())
        ext = format  # "md" | "pdf" | "docx"
        out_dir = Path.home() / ".piighost" / "exports"
        out_dir.mkdir(parents=True, exist_ok=True)
        output_path = str(out_dir / f"{project}-{doctype}-{ts}.{ext}")

    # Containment: the resolved output must live under ~/.piighost/. This
    # blocks MCP callers from writing arbitrary files via ../ traversal or
    # absolute paths.
    exports_root = (Path.home() / ".piighost").resolve()
    out_resolved = Path(output_path).resolve()
    try:
        out_resolved.relative_to(exports_root)
    except ValueError as e:
        raise ValueError(
            f"output_path must be under {exports_root}; got {out_resolved}",
        ) from e

    out = out_resolved
    out.parent.mkdir(parents=True, exist_ok=True)

    # Always render the Markdown body — both PDF and DOCX paths build on it.
    md_body = _render_md(data, profile=profile, doctype=doctype)

    if format == "md":
        out.write_text(md_body, encoding="utf-8")
        size = out.stat().st_size
    elif format == "pdf":
        size = _render_pdf(md_body, out)
    else:  # docx
        size = _render_docx(md_body, out)

    return {
        "path": str(out),
        "format": format,
        "size_bytes": size,
        "rendered_at": int(time.time()),
    }
